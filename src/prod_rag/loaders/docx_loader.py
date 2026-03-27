from __future__ import annotations

from pathlib import Path

from docx import Document

from prod_rag.loaders.base import BaseLoader
from prod_rag.models import LoadedDocument


class DocxLoader(BaseLoader):
    def load(self, path: str | Path) -> LoadedDocument:
        p = Path(path)
        doc = Document(str(p))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return LoadedDocument(document_id=p.stem, source_path=str(p), text=text, metadata={"type": "docx"})
